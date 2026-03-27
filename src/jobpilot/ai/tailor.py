"""
Resume tailor — customize resume for a specific job posting.

Methodology adapted from Proficiently skill suite:
- Tailoring philosophy: reorganize and reframe real experience
- Strict accuracy: never fabricate, never inflate, never assume
- Anti-AI writing: no emdashes, no filler, Flesch >90
- Quality checks: 7-second scan test, every bullet traces to source
"""

from __future__ import annotations

import logging
from pathlib import Path

import yaml

from jobpilot import config
from jobpilot.models import Job, Profile

logger = logging.getLogger(__name__)

RESUME_CONFIG_PATH = config.DATA_DIR / "resume_config.yaml"

# ---------------------------------------------------------------------------
# Prompt — full Proficiently methodology
# ---------------------------------------------------------------------------

TAILOR_PROMPT = """\
You are an expert resume writer who creates tailored resumes that make \
candidates the obvious choice for a specific role.

## Tailoring Philosophy

The goal is NOT to make the candidate look like someone they're not. \
The goal is to reorganize and reframe their real experience so the hiring \
manager immediately sees the fit.

A hiring manager spends ~7 seconds on a first-pass resume scan. In those \
7 seconds, they should see:
- A summary that speaks directly to their open role
- The most relevant experience front and center
- Keywords that match their job description
- Evidence of operating at the right level

## Resume Sections

### Summary (2-3 sentences)
- Reference the type of role and industry directly
- Lead with the most relevant credential (years of experience, biggest \
result, most relevant company)
- Include 2-3 keywords from the job posting naturally
- End with what makes this candidate distinctive
- Each sentence under 20 words. No run-on sentences.

Good: "Economics graduate with hands-on Python and ML experience from \
research and independent data projects. Built automated data pipelines \
and predictive models achieving 86% precision. Brings strong quantitative \
foundation in causal inference and econometric modeling."

Bad: "Results-driven leader with a passion for data and a track record \
of success in fast-paced environments."

### Experience

For each role, select and order bullets by relevance to the target job:

Bullet formula: [Action verb] + [what you did] + [how/at what scale] + \
[measurable result]

- First 2 bullets of each role = most relevant to the target job
- Metrics in every bullet where possible
- Mirror job posting language where authentic
- Remove irrelevant bullets rather than leaving noise
- Each bullet should start with a strong action verb
- Each bullet should show: what you did, how you did it, what the impact was

Bullet count per role:
- Current/most recent role: 4-6 bullets
- Previous roles: 3-4 bullets
- Older roles: 2-3 bullets

### Skills Section
- Reorganize to lead with skills the job posting emphasizes
- Group into categories that match the job's framing
- Remove skills that are irrelevant noise for this specific role
- ONLY use skills from the VERIFIED SKILLS list below

### Education
- Same as original. Do not change.

## Level Calibration

For senior IC / manager roles:
- Emphasize: hands-on expertise, technical depth, mentorship, direct impact
- Show collaboration and influence without authority
- Concrete deliverables and project outcomes

For intern / entry-level roles:
- Emphasize: relevant coursework, project experience, technical skills, \
eagerness to learn
- Show initiative through independent projects
- Academic achievements and research experience

## Writing Rules (CRITICAL — target Flesch score above 90)

- Write like a sharp professional, not a language model. Short sentences. \
Plain words.
- Every sentence gets one idea. If a sentence has "and" connecting two \
unrelated clauses, split it.
- Never use emdashes (—). Use commas, periods, colons, semicolons, or \
parentheses instead.
- Vary sentence structure. Not every bullet should follow the exact same \
pattern.
- No preamble clauses. Bad: "Leveraging deep expertise in data, led..." \
Good: "Led..."
- No stacking adjectives. Bad: "cross-functional, data-driven, \
customer-centric approach". Pick one.
- No filler phrases: "demonstrating ability to", "showcasing expertise in", \
"with a track record of", "needed to drive", "spanning", "leveraging", \
"utilizing"
- No compound noun piles: "AI-driven product opportunity identification \
and execution" — just say what you did
- Summaries must be 2-3 SHORT sentences. Each sentence under 20 words.

## Strict Accuracy Rules (CRITICAL — non-negotiable)

- ONLY use information explicitly provided in the resume or verified skills. \
NEVER fill gaps with assumptions.
- Never assume business model: Don't label a company as B2B, B2C, SaaS, etc. \
unless explicitly stated.
- Never inflate scope: If the resume says "revenue targets," don't write \
"P&L ownership." If it says "data analysis," don't write "data engineering."
- Never add cross-functional partners not mentioned.
- When reframing, only reframe what exists. You can reorder bullets, change \
wording, and mirror job posting language, but every claim must trace back to \
a specific fact from the source materials.
- If something is ambiguous, use conservative language or omit it. Better to \
understate than overstate.
- Job titles and dates are UNCHANGED from original.
- Company names are UNCHANGED from original.

## What NOT to Do

- Don't fabricate experience or skills the candidate doesn't have
- Don't use generic buzzwords that aren't backed by specific experience
- Don't make the resume longer than 2 pages
- Don't change job titles or dates
- Don't remove roles (gaps look suspicious)
- Don't use Markdown formatting (no #, **, -, etc.) — output plain text
- Don't add emoji or decorative symbols
- Don't add Chinese annotations to English skills or vice versa

## Self-Critique Checklist (apply BEFORE outputting)

Go sentence by sentence and fix AI-sounding writing:

1. Is this sentence doing too much? If it has more than one comma-separated \
clause, split it.
2. Would a real person say this? If it sounds like LinkedIn or ChatGPT, \
rewrite it.
3. Is there filler? Cut any phrase that doesn't add information.
4. Are there stacked buzzwords? Pick the one that matters and give a \
concrete example.
5. Is the summary under control? Max 3 sentences. Each under 20 words.

Common AI patterns to kill:
- "I combine X with Y, Z, and the W needed to..." -> Split into separate \
statements
- "...demonstrating [abstract quality]" -> Delete or replace with actual result
- "...spanning [long list]" -> Pick the most relevant 1-2 items
- "Led [action], [action], and [action] across [scope]" -> One action per bullet
- Any bullet over 2 lines is probably trying to do too much -> split it
- Gerund clauses tacked onto the end: "...delivering X while maintaining Y" \
-> Two sentences

Test: After writing, re-read the summary and first 3 bullets. If any sentence \
takes more than one breath to read out loud, it's too long. Shorten it.

## Quality Checks (verify before returning)

- [ ] Summary references the specific role/industry
- [ ] Most relevant experience appears in the first 2 bullets of each role
- [ ] Metrics appear in at least 60% of bullets
- [ ] Keywords from the job posting appear naturally throughout
- [ ] No fabricated experience or inflated titles
- [ ] Job titles and dates are unchanged from original
- [ ] Action verbs are varied (not all "Led" or "Managed")
- [ ] Level of language matches the role's seniority
- [ ] A hiring manager scanning for 7 seconds would see the fit
- [ ] Every bullet traces back to a specific fact from the resume
- [ ] Business model, scope, and responsibilities match what's actually stated
- [ ] Output is plain text (no Markdown formatting)

---

## CANDIDATE'S ORIGINAL RESUME
{resume_text}

## CANDIDATE'S VERIFIED SKILLS (do NOT add skills beyond this list)
{skills_text}

## TARGET JOB
- Title: {job_title}
- Company: {company}
- Job Description:
{jd_text}

## CONTACT INFORMATION (include in resume header)
{contact_info}

## OUTPUT
Output the tailored resume as plain text. Follow the candidate's original \
resume structure. Only output the resume content, no explanations or \
commentary. Preserve the original language (English resume -> English output, \
Chinese resume -> Chinese output)."""


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def _load_skills_config() -> dict:
    """Load skills from resume_config.yaml if it exists."""
    if not RESUME_CONFIG_PATH.exists():
        return {}
    with open(RESUME_CONFIG_PATH, encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


def _format_contact_info(skills_config: dict) -> str:
    """Format contact information from resume_config.yaml for the prompt."""
    parts: list[str] = []
    name = skills_config.get("name", "")
    if name:
        parts.append(f"Name: {name}")
    email = skills_config.get("email", "")
    if email:
        parts.append(f"Email: {email}")
    phone = skills_config.get("phone", "")
    if phone:
        parts.append(f"Phone: {phone}")
    location = skills_config.get("location", "")
    if location:
        parts.append(f"Location: {location}")
    github = skills_config.get("github", "")
    if github:
        parts.append(f"GitHub: {github}")
    return "\n".join(parts) if parts else "(no contact info configured)"


def _format_skills_text(skills_config: dict, profile: Profile) -> str:
    """Format skills for the prompt. Prefer config file, fall back to profile."""
    skill_groups = skills_config.get("skills", [])
    if skill_groups:
        lines = []
        for group in skill_groups:
            category = group.get("category", "")
            items = group.get("items", [])
            lines.append(f"{category}: {', '.join(items)}")
        return "\n".join(lines)

    # Fallback to profile structured data
    s = profile.structured
    skills = s.get("skills", {})
    if isinstance(skills, dict):
        lines = []
        for category, items in skills.items():
            if isinstance(items, list) and items:
                lines.append(f"{category}: {', '.join(items)}")
        return "\n".join(lines)
    return ""


# ---------------------------------------------------------------------------
# Core API
# ---------------------------------------------------------------------------

def tailor_resume(profile: Profile, job: Job) -> str:
    """Generate a tailored resume for a specific job.

    Uses the full Proficiently methodology:
    - Reorganize and reframe real experience for the target role
    - Strict accuracy rules (never fabricate, never inflate)
    - Anti-AI writing rules (no emdashes, no filler, Flesch >90)
    - Quality checks (7-second scan test, every bullet traces to source)
    """
    skills_config = _load_skills_config()
    skills_text = _format_skills_text(skills_config, profile)
    contact_info = _format_contact_info(skills_config)

    prompt = TAILOR_PROMPT.format(
        resume_text=profile.raw_text,
        skills_text=skills_text,
        contact_info=contact_info,
        job_title=job.title,
        company=job.company,
        jd_text=job.jd_text,
    )

    if not config.ANTHROPIC_API_KEY:
        logger.warning("ANTHROPIC_API_KEY not set. Cannot tailor resume.")
        return _basic_tailor(profile, job)

    import anthropic

    try:
        client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
        message = client.messages.create(
            model=config.ANTHROPIC_MODEL,
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        return message.content[0].text.strip()
    except (anthropic.BadRequestError, anthropic.AuthenticationError) as e:
        logger.warning("API call failed (%s), falling back to prompt export.", e)
        return _basic_tailor(profile, job)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace text in a paragraph while preserving all run formatting."""
    if not paragraph.runs:
        return
    # Put all new text in the first run, clear the rest
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _patch_docx(source_docx: Path, tailored_text: str, output_path: Path) -> None:
    """Copy original docx and replace text content, preserving all formatting.

    Uses difflib SequenceMatcher to align original paragraphs with tailored
    lines, then only replaces paragraphs where text actually changed.
    Structural paragraphs (bold titles, tab-aligned dates) are never touched.
    """
    import shutil
    from difflib import SequenceMatcher
    from docx import Document

    shutil.copy2(source_docx, output_path)
    doc = Document(str(output_path))

    # Build indexed list of non-empty paragraphs
    orig_entries = []  # (para_index, stripped_text)
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            orig_entries.append((i, t))

    tailored_lines = [l.strip() for l in tailored_text.split("\n") if l.strip()]

    orig_texts = [e[1] for e in orig_entries]

    # Align using SequenceMatcher
    matcher = SequenceMatcher(None, orig_texts, tailored_lines, autojunk=False)

    # Paragraphs we should NOT touch: bold title lines, section headers, tab-aligned
    def _is_structural(para) -> bool:
        text = para.text.strip()
        if not text:
            return True
        # All-caps section headers
        if text.isupper() and len(text) > 2:
            return True
        # Bold title/institution lines
        has_non_empty_run = any(r.text.strip() for r in para.runs)
        if has_non_empty_run and all(r.bold for r in para.runs if r.text.strip()):
            return True
        # Tab-aligned role/date lines (but not skills lines)
        if "\t" in text:
            return True
        return False

    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == "equal":
            continue
        if op == "replace":
            # Replace matched pairs 1:1
            for k in range(min(i2 - i1, j2 - j1)):
                para_idx = orig_entries[i1 + k][0]
                para = doc.paragraphs[para_idx]
                if not _is_structural(para):
                    _replace_paragraph_text(para, tailored_lines[j1 + k])

    # Inject GitHub link if configured
    cfg = _load_skills_config()
    github_url = cfg.get("github", "")
    if github_url:
        _inject_github_link(doc, github_url)

    doc.save(str(output_path))


def _inject_github_link(doc, github_url: str) -> None:
    """Inject GitHub link into the contact info line of a docx document.

    Scans the first 5 paragraphs for a line containing email/@/phone,
    then appends ' | GitHub: <url>' if not already present.
    """
    for para in doc.paragraphs[:5]:
        text = para.text.strip()
        if not text:
            continue
        # Look for contact line (contains @ or phone-like pattern)
        has_contact = ("@" in text or
                       any(c.isdigit() for c in text if len(text) < 200))
        if has_contact and "github" not in text.lower():
            # Append to the last non-empty run
            for run in reversed(para.runs):
                if run.text.strip():
                    run.text = run.text + f" | GitHub: {github_url}"
                    return
    # If no contact line found, don't inject


def _find_source_docx() -> Path | None:
    """Find the original resume docx from resume_config.yaml."""
    cfg = _load_skills_config()
    source = cfg.get("resume_source", "")
    if source:
        p = Path(source)
        if p.exists():
            return p
    return None


def save_tailored_resume(
    profile: Profile,
    job: Job,
    output_dir: Path | None = None,
) -> Path:
    """Generate and save a tailored resume as Word document.

    If original .docx is available, patches it to preserve formatting.
    Otherwise falls back to plain text .txt.

    Returns:
        Path to the saved file
    """
    output_dir = output_dir or config.TAILORED_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    content = tailor_resume(profile, job)
    safe_company = job.company.replace("/", "_")[:20]
    safe_title = job.title.replace("/", "_")[:30]

    source_docx = _find_source_docx()
    if source_docx:
        filename = f"{safe_company}_{safe_title}.docx"
        output_path = output_dir / filename
        _patch_docx(source_docx, content, output_path)
    else:
        filename = f"{safe_company}_{safe_title}.txt"
        output_path = output_dir / filename
        output_path.write_text(content, encoding="utf-8")

    logger.info("Tailored resume saved to %s", output_path)
    return output_path


def _basic_tailor(profile: Profile, job: Job) -> str:
    """Fallback when AI is not available: export the full prompt for manual use.

    Saves the complete tailor prompt to data/tailored/{company}_{title}_prompt.txt
    so the user can paste it into Claude.ai or another LLM.
    Returns the original resume text (for backward compatibility).
    """
    skills_config = _load_skills_config()
    skills_text = _format_skills_text(skills_config, profile)
    contact_info = _format_contact_info(skills_config)

    prompt = TAILOR_PROMPT.format(
        resume_text=profile.raw_text,
        skills_text=skills_text,
        contact_info=contact_info,
        job_title=job.title,
        company=job.company,
        jd_text=job.jd_text,
    )

    # Save prompt file
    prompt_dir = config.TAILORED_DIR
    prompt_dir.mkdir(parents=True, exist_ok=True)
    safe_company = job.company.replace("/", "_")[:20]
    safe_title = job.title.replace("/", "_")[:30]
    prompt_path = prompt_dir / f"{safe_company}_{safe_title}_prompt.txt"
    prompt_path.write_text(prompt, encoding="utf-8")

    logger.info("Tailor prompt exported to %s", prompt_path)
    print(f"\n[无 API 模式] 已导出完整 prompt → {prompt_path}")
    print("请将 prompt 内容粘贴到 Claude.ai，获取定制简历文本后运行：")
    print(f"  jobpilot tailor {job.job_id} --from-text <output_file>\n")

    # Return original text for backward compatibility
    target_line = f"TARGET: {job.title} @ {job.company}\n\n"
    return target_line + profile.raw_text


def tailor_from_text(
    text: str,
    job: Job,
    output_dir: Path | None = None,
) -> Path:
    """Import externally generated tailored text and patch into docx.

    This allows users to paste the output from Claude.ai back into the
    pipeline to get a properly formatted .docx file.

    Args:
        text: The tailored resume text (from Claude.ai or other LLM)
        job: The target job
        output_dir: Where to save; defaults to config.TAILORED_DIR

    Returns:
        Path to the saved file (.docx or .txt)
    """
    output_dir = output_dir or config.TAILORED_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_company = job.company.replace("/", "_")[:20]
    safe_title = job.title.replace("/", "_")[:30]

    source_docx = _find_source_docx()
    if source_docx:
        filename = f"{safe_company}_{safe_title}.docx"
        output_path = output_dir / filename
        _patch_docx(source_docx, text, output_path)
    else:
        filename = f"{safe_company}_{safe_title}.txt"
        output_path = output_dir / filename
        output_path.write_text(text, encoding="utf-8")

    logger.info("Tailored resume (from text) saved to %s", output_path)
    return output_path
